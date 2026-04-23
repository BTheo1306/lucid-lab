"""
Generates the bot tuning questionnaire as a formatted .docx file.
Run: python3 scripts/generate-questionnaire-docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── helpers ──────────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text="", bold=False, italic=False, color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_question(doc, number, question, hint=None):
    """Bold question line, optional grey hint below, then answer box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{number}  {question}")
    run.bold = True
    run.font.size = Pt(10.5)

    if hint:
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(0)
        ph.paragraph_format.space_after = Pt(2)
        rh = ph.add_run(hint)
        rh.italic = True
        rh.font.size = Pt(9)
        rh.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    add_answer_box(doc)

def add_answer_box(doc, lines=2):
    """Shaded grey box for written answers."""
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.right_indent = Inches(0.1)
        shade_paragraph(p, "F0F0F0")
        p.add_run(" " * 80)
        p.runs[0].font.size = Pt(10)
    # small gap after box
    gap = doc.add_paragraph()
    gap.paragraph_format.space_after = Pt(6)

def add_checkboxes(doc, options, allow_other=False):
    """List of checkbox items."""
    for opt in options:
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"☐  {opt}")
        run.font.size = Pt(10)
    if allow_other:
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run("☐  Autre : ")
        run.font.size = Pt(10)
        run2 = p.add_run("_" * 40)
        run2.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def shade_paragraph(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)

def add_section_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("─" * 90)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default body font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# ── Cover ─────────────────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(24)
p_title.paragraph_format.space_after = Pt(4)
r = p_title.add_run("Bot Tuning Questionnaire")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(6)
r2 = p_sub.add_run("Lucid-Lab  ·  Chat Widget")
r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p_inst = doc.add_paragraph()
p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_inst.paragraph_format.space_after = Pt(32)
r3 = p_inst.add_run(
    "Remplissez les champs grisés. Aucun format imposé — mots-clés, phrases ou exemples.\n"
    "Les exemples d'échanges (section 2) sont la partie la plus utile."
)
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
r3.italic = True

add_section_divider(doc)

# ── Section 1 ─────────────────────────────────────────────────────────────────
set_heading(doc, "1. Identité & voix du bot", level=1)

add_question(doc, "1.1", "Nom du bot",
             "Il en a un ? (ex. « Lucid », « Lab », « Clara », ou aucun)")

add_question(doc, "1.2", "Persona en une phrase",
             "Qui joue-t-il ? (ex. « un ingénieur solutions en interne », « un consultant senior franc du collier »)")

add_para(doc, "1.3  Ton", bold=True)
add_checkboxes(doc, [
    "Professionnel / corporate",
    "Chaleureux & conversationnel (défaut actuel)",
    "Direct / sans fioritures",
    "Avec une pointe d'humour",
], allow_other=True)

add_para(doc, "1.4  Tutoiement ou vouvoiement (FR) ?", bold=True)
add_checkboxes(doc, ["Tutoiement", "Vouvoiement", "Laisser le bot s'adapter selon le registre du visiteur"])

add_para(doc, "1.5  Emojis ?", bold=True)
add_checkboxes(doc, ["Oui, librement", "Oui, mais avec modération", "Non"])

add_para(doc, "1.6  Mise en forme des réponses", bold=True)
add_checkboxes(doc, [
    "Riche — titres, tableaux, listes (défaut actuel)",
    "Paragraphes conversationnels, pas de markdown visible",
    "Mix selon la complexité de la question",
])

add_para(doc, "1.7  Longueur des réponses", bold=True)
add_checkboxes(doc, [
    "Courte (2-4 phrases max)",
    "Moyenne (défaut actuel)",
    "Longue / détaillée si nécessaire",
])

add_section_divider(doc)

# ── Section 2 ─────────────────────────────────────────────────────────────────
set_heading(doc, "2. Trois exemples d'échanges  ← section la plus importante", level=1)

add_para(doc, "Inventez ou racontez 3 échanges. Gauche = visiteur, droite = réponse idéale. "
         "La qualité du résultat dépend directement de ces exemples.", italic=True,
         color=(0x77, 0x77, 0x77))

for label, desc in [
    ("A", "Question classique d'un prospect"),
    ("B", "Question sur les prix / budget"),
    ("C", "Question hors-sujet, délicate ou piège (ex. « l'IA va-t-elle remplacer mon équipe ? »)"),
]:
    add_para(doc, f"Exemple {label}  —  {desc}", bold=True, space_before=8)
    add_para(doc, "Visiteur :", italic=True)
    add_answer_box(doc, lines=2)
    add_para(doc, "Réponse idéale :", italic=True)
    add_answer_box(doc, lines=4)

add_section_divider(doc)

# ── Section 3 ─────────────────────────────────────────────────────────────────
set_heading(doc, "3. Ce que le bot DOIT faire", level=1)

add_para(doc, "3.1  Objectif principal — numérotez par priorité (1 = plus important)", bold=True)
add_checkboxes(doc, [
    "Réserver un appel découverte (TidyCal)",
    "Capturer les coordonnées du lead (email, entreprise, besoin)",
    "Éduquer / répondre aux questions",
    "Qualifier le fit avant de proposer un appel",
])

add_question(doc, "3.2", "Quand doit-il pousser vers la réservation ?",
             "Ex. « après 2 échanges », « dès qu'il identifie un pain point », « seulement si demandé »")

add_question(doc, "3.3", "Infos à collecter avant de proposer un appel",
             "Cochez ou listez : nom, email, entreprise, taille équipe, budget, délai, besoin…")

add_question(doc, "3.4", "Lien de réservation",
             "Actuel : https://tidycal.com/lucid-lab/audit-flash-30-minutes  —  À modifier ?")

add_section_divider(doc)

# ── Section 4 ─────────────────────────────────────────────────────────────────
set_heading(doc, "4. Ce que le bot NE DOIT PAS faire", level=1)

add_question(doc, "4.1", "Sujets à refuser ou esquiver",
             "Ex. conseil juridique, garanties de ROI, comparaisons de concurrents, politique…")

add_question(doc, "4.2", "Affirmations à ne jamais faire",
             "Ex. « ROI garanti », pourcentages précis, délais, noms de clients sans permission")

add_question(doc, "4.3", "Règles absolues",
             "Ex. « ne jamais donner de fourchette de prix », « toujours recommander un audit avant devis »")

add_section_divider(doc)

# ── Section 5 ─────────────────────────────────────────────────────────────────
set_heading(doc, "5. Contenu & positionnement", level=1)

add_question(doc, "5.1", "Pitch en une phrase",
             "Si vous deviez réécrire le positionnement de Lucid-Lab en une seule phrase")

add_question(doc, "5.2", "Client idéal",
             "Secteur, taille, stade de maturité IA, rôle de l'acheteur")

add_question(doc, "5.3", "Qui n'est PAS un bon fit",
             "Le bot doit poliment rediriger : ex. solo freelances, TPE < 10 salariés, budget < X€")

add_para(doc, "5.4  Top 3 services à mettre en avant", bold=True)
for i in range(1, 4):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"{i}.  ")
    run.font.size = Pt(10.5)
    shade_paragraph(p, "F0F0F0")
    p.add_run(" " * 60)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_question(doc, "5.5", "Différenciateurs vs consultants / agences classiques",
             "2-3 bullets sur ce qui vous distingue")

add_question(doc, "5.6", "Preuves sociales / cas d'usage que le bot peut mentionner",
             "Noms réels, secteurs, résultats — ou « garder générique » ?")

add_para(doc, "5.7  Quand quelqu'un demande « combien ça coûte ? »", bold=True)
add_checkboxes(doc, [
    "Donner une fourchette (laquelle ?  ___________________________)",
    "Rediriger vers l'appel d'audit",
    "Estimer selon le type de projet",
    "Expliquer la valeur avant de parler prix",
], allow_other=False)

add_section_divider(doc)

# ── Section 6 ─────────────────────────────────────────────────────────────────
set_heading(doc, "6. Langues & localisation", level=1)

add_para(doc, "6.1  Langues supportées", bold=True)
add_checkboxes(doc, ["FR + EN (défaut actuel)", "FR uniquement", "Ajouter d'autres langues (lesquelles ?)"], allow_other=True)

add_para(doc, "6.2  Langue par défaut si ambiguë", bold=True)
add_checkboxes(doc, ["Français", "Anglais", "Laisser le bot détecter"])

add_question(doc, "6.3", "Vocabulaire préféré",
             "Ex. « solutions IA » vs « intelligence artificielle », « clients » vs « partenaires », « audit » vs « diagnostic »")

add_section_divider(doc)

# ── Section 7 ─────────────────────────────────────────────────────────────────
set_heading(doc, "7. Cas limites & escalade", level=1)

add_question(doc, "7.1", "Quand le visiteur est frustré ou bloqué",
             "Ex. proposer un transfert humain, partager un email direct, s'excuser et clore")

add_question(doc, "7.2", "Contact direct à partager si besoin",
             "Email, téléphone — ou « jamais, toujours renvoyer vers la réservation »")

add_question(doc, "7.3", "Requêtes urgentes ou sensibles",
             "Ex. « il me faut ça pour vendredi », « mon DG est en attente »")

add_question(doc, "7.4", "Quelqu'un demande un livrable gratuit",
             "Ex. « écris-moi un prompt », « génère un plan d'automatisation » — aider un peu / esquiver / refuser ?")

add_section_divider(doc)

# ── Section 8 ─────────────────────────────────────────────────────────────────
set_heading(doc, "8. Widget & UX  (optionnel)", level=1)

add_question(doc, "8.1", "Premier message à l'ouverture du chat",
             "Laisser le bot décider, ou imposer un texte spécifique ?")

add_para(doc, "8.2  Boutons de démarrage rapide (chips)", bold=True)
add_para(doc, "3 boutons affichés dès l'ouverture. Exemples : « Vos services », « Tarifs », « Prendre RDV »", italic=True, color=(0x77,0x77,0x77))
for i in range(1, 4):
    p = doc.add_paragraph(style="List Paragraph")
    shade_paragraph(p, "F0F0F0")
    p.add_run(f"Bouton {i} :  " + " " * 40)
    p.runs[0].font.size = Pt(10.5)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_para(doc, "8.3  Apparition de la bulle", bold=True)
add_checkboxes(doc, [
    "Immédiatement",
    "Après X secondes (combien ?  ___)",
    "Au scroll (50 % de la page)",
    "Uniquement sur certaines pages",
])

add_section_divider(doc)

# ── Section 9 ─────────────────────────────────────────────────────────────────
set_heading(doc, "9. Notes libres", level=1)
add_para(doc, "Tout ce qui n'entre dans aucune case ci-dessus : brief de marque, ton de la comm existante, "
         "liens, captures d'écran à joindre, concurrents à connaître…", italic=True, color=(0x77,0x77,0x77))
add_answer_box(doc, lines=8)

# ── Footer note ──────────────────────────────────────────────────────────────
add_section_divider(doc)
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run(
    "Une fois complété, renvoyez ce fichier — les réponses seront traduites en :\n"
    "system prompt mis à jour · nouvelles entrées knowledge base · règles de guardrails · ajustements UI"
)
r_f.font.size = Pt(9)
r_f.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r_f.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
output = "bot-tuning-questionnaire.docx"
doc.save(output)
print(f"Saved: {output}")
