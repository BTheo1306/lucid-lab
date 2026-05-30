#!/usr/bin/env python3
"""
Lucid-Lab — NDA Word document generator (python-docx)
Uses the BDC source template as the base so the NDA shares the exact same
layout, logo, fonts and styles as the Proposition / Contrat de prestation.

Usage:
    python3 scripts/generate_nda_word.py
"""
from __future__ import annotations

import shutil
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
REPO_ROOT   = Path(__file__).parent.parent
OUT_DIR     = REPO_ROOT / "docs" / "legal-templates" / "generated"
BDC_SOURCE  = REPO_ROOT / "docs" / "legal-templates" / "source" / "03_Proposition-valant-Bon-de-commande-modele-logo.docx"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# Colours — same as BDC
# ---------------------------------------------------------------------------
INK   = RGBColor(0x00, 0x00, 0x00)
MID   = RGBColor(0x66, 0x66, 0x66)
EMBER = RGBColor(0xC8, 0x5E, 0x1A)

# ---------------------------------------------------------------------------
# Base: clone BDC source, clear body, update header subtitle
# ---------------------------------------------------------------------------

def _clone_bdc(tmp_path: Path) -> Document:
    """Copy BDC source to tmp_path and open it."""
    shutil.copy(str(BDC_SOURCE), str(tmp_path))
    return Document(str(tmp_path))


def _clear_body(doc: Document) -> None:
    """Remove all body paragraphs and tables, leaving section properties."""
    body_el = doc.element.body
    # Keep only sectPr (last element)
    children = list(body_el)
    sect_pr = body_el.find(qn("w:sectPr"))
    for child in children:
        if child is not sect_pr:
            body_el.remove(child)


def _update_header_subtitle(doc: Document, subtitle: str) -> None:
    """Replace the header label text (3rd run in header para) with subtitle."""
    sec = doc.sections[0]
    header = sec.header
    if not header.paragraphs:
        return
    p = header.paragraphs[0]
    # Runs: [0]=Lucid-Lab wordmark, [1]=spacing, [2]=subtitle, ...
    # Find the subtitle run (non-empty, non-Lucid-Lab, not just spaces)
    for run in p.runs:
        if run.text and run.text.strip() and "Lucid-Lab" not in run.text and run.text.strip() not in ("", "    "):
            run.text = subtitle
            break


def _update_footer(doc: Document, text: str) -> None:
    sec = doc.sections[0]
    footer = sec.footer
    if footer.paragraphs:
        p = footer.paragraphs[0]
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = text
        else:
            r = p.add_run(text)
            r.font.size = Pt(7.5)


# ---------------------------------------------------------------------------
# Content helpers — match BDC run style exactly
# ---------------------------------------------------------------------------

def _add_title(doc: Document, text: str) -> None:
    """22pt Syne Bold title — same as BDC."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(12)
    r = p.add_run(text)
    r.font.name = "Syne"
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = INK


def _add_meta(doc: Document, text: str) -> None:
    """Small grey meta line (ref, date) — same as BDC."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(14)
    r = p.add_run(text)
    r.font.color.rgb = MID


def _add_intro(doc: Document, text: str) -> None:
    """Italic intro paragraph in grey."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(14)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.color.rgb = MID
    r.font.italic = True


def _add_section(doc: Document, text: str) -> None:
    """15pt Syne Bold section heading — identical to BDC sections."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Syne"
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = INK


def _add_body(doc: Document, text: str, space_after: float = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)


def _add_body_mixed(doc: Document, parts: list[tuple[str, bool]], space_after: float = 4) -> None:
    """Body paragraph with mixed bold/normal fragments."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in parts:
        r = p.add_run(text)
        r.font.bold = bold


def _add_bullet(doc: Document, text: str) -> None:
    """Bullet using same • prefix style as BDC."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.add_run(f"\u2022  {text}")


def _add_party_box(doc: Document, label: str, lines: list[str]) -> None:
    """Bordered box for a party — matches BDC party block style."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table1"
    cell = tbl.cell(0, 0)
    # light grey shading
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F8F8F8")
    tcPr.append(shd)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    lp = cell.add_paragraph()
    lr = lp.add_run(label)
    lr.font.bold = True
    lr.font.size = Pt(9.5)
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after  = Pt(2)
    for line in lines:
        p  = cell.add_paragraph()
        p.add_run(line)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
    # spacer after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def _add_penalty_box(doc: Document, intro: str, amount: str, caveat: str) -> None:
    """Orange-accented penalty box — matches BDC penalty clause style."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table1"
    cell = tbl.cell(0, 0)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # shading
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "FDF6F2")
    tcPr.append(shd)
    # left border in Ember
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        if side == "left":
            bd.set(qn("w:val"),  "thick")
            bd.set(qn("w:sz"),   "18")
            bd.set(qn("w:color"), "C85E1A")
        else:
            bd.set(qn("w:val"),  "single")
            bd.set(qn("w:sz"),   "4")
            bd.set(qn("w:color"), "D0D0D0")
        tcBorders.append(bd)
    tcPr.append(tcBorders)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for text, bold, size_pt, color_rgb in [
        (intro,  False, 9,    INK),
        (amount, True,  12,   EMBER),
        (caveat, False, 9,    INK),
    ]:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        r.font.bold = bold
        r.font.size = Pt(size_pt)
        r.font.color.rgb = color_rgb
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def _add_context_box(doc: Document, text: str) -> None:
    """Indented context/mission box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table1"
    cell = tbl.cell(0, 0)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.italic = True
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def _add_signature_table(doc: Document,
                          party_a_label: str, party_b_label: str,
                          signer_name: str) -> None:
    """Signature table — same structure as BDC."""
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Table1"
    rows_data = [
        [party_a_label,  party_b_label],
        ["Date : ___________________", "Date : ___________________"],
        ["Lieu : ___________________", "Lieu : ___________________"],
        [f"Nom et signature précédés de\n« Lu et approuvé, bon pour accord »\n\n\n_________________________",
         f"Nom et signature précédés de\n« Lu et approuvé, bon pour accord »\n\n\n_________________________"],
    ]
    for ri, row_data in enumerate(rows_data):
        for ci, text in enumerate(row_data):
            cell = tbl.cell(ri, ci)
            cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
            p  = cell.add_paragraph()
            rn = p.add_run(text)
            rn.font.bold = (ri == 0)
            rn.font.size = Pt(9.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(14 if ri == 3 else 2)

# ---------------------------------------------------------------------------
# NDA content builder — shared logic, lang switch controls wording
# ---------------------------------------------------------------------------

def _build_nda(
    lang: str,          # "fr" | "en"
    doc_ref: str,
    issued_date: str,
    client_name: str,
    client_address: str,
    client_reg: str,
    signer_name: str,
    signer_title: str,
    signer_email: str,
    mission_context: str,
    nda_duration: str,
    out_path: Path,
) -> Path:
    # Clone BDC source to get exact layout, logo, styles
    tmp = out_path.with_suffix(".tmp.docx")
    doc = _clone_bdc(tmp)
    _clear_body(doc)

    is_fr = (lang == "fr")

    # --- header subtitle
    _update_header_subtitle(
        doc,
        "Accord de Non-Divulgation" if is_fr else "Non-Disclosure Agreement"
    )

    # --- footer
    _update_footer(
        doc,
        "Lucid-Lab · SAS au capital de 999 € · RCS Paris 104 672 050 · TVA FR 02 104 672 050 · 47 rue Vivienne, 75002 Paris"
        + ("" if is_fr else " · France")
    )

    # ---- Title
    _add_title(doc, "ACCORD DE NON-DIVULGATION" if is_fr else "NON-DISCLOSURE AGREEMENT")

    # ---- Meta line
    if is_fr:
        _add_meta(doc, f"Référence : {doc_ref}   ·   Date d'émission : {issued_date}   ·   Document confidentiel")
    else:
        _add_meta(doc, f"Reference: {doc_ref}   ·   Date: {issued_date}   ·   Confidential document")

    # ---- Intro
    if is_fr:
        _add_intro(doc,
            "Le présent Accord de Non-Divulgation (« l'Accord ») encadre les échanges d'informations confidentielles "
            "dans le cadre des discussions précontractuelles entre les Parties. Il est conclu à titre mutuel : "
            "chaque Partie peut être à la fois Partie Divulgatrice et Partie Réceptrice.")
    else:
        _add_intro(doc,
            "This Non-Disclosure Agreement (the \"Agreement\" or \"NDA\") governs the exchange of confidential information "
            "in the context of pre-contractual discussions between the Parties. It is entered into on a mutual basis: "
            "each Party may act as both a Disclosing Party and a Receiving Party.")

    # ---- Parties
    _add_section(doc, "Parties" if is_fr else "Parties")
    if is_fr:
        _add_party_box(doc, "Partie A — Lucid-Lab (le Prestataire)", [
            "Lucid-Lab, Société par actions simplifiée (SAS) au capital de 999 euros,",
            "Siège social : 47 rue Vivienne, 75002 Paris, France",
            "RCS Paris n° 104 672 050  ·  TVA FR 02 104 672 050",
            "Représentée par son Président, Periscope-X SARL,",
            "elle-même représentée par M. Anthony Poirier, dûment habilité.",
        ])
        _add_party_box(doc, "Partie B — Le Client / Destinataire", [
            f"Société : {client_name}",
            f"Adresse : {client_address}",
            f"SIRET / N° d'immatriculation : {client_reg}",
            f"Représentée par : {signer_name}, en qualité de {signer_title}",
            f"E-mail : {signer_email}",
        ])
        _add_body(doc, "Ensemble dénommées les « Parties ».")
    else:
        _add_party_box(doc, "Party A — Lucid-Lab (the Service Provider)", [
            "Lucid-Lab, Société par actions simplifiée (SAS), registered capital €999,",
            "Registered office: 47 rue Vivienne, 75002 Paris, France",
            "RCS Paris no. 104 672 050  ·  VAT FR 02 104 672 050",
            "Represented by its Chairman, Periscope-X SARL,",
            "itself represented by Mr Anthony Poirier.",
        ])
        _add_party_box(doc, "Party B — The Client", [
            f"Company: {client_name}",
            f"Address: {client_address}",
            f"Registration no.: {client_reg}",
            f"Represented by: {signer_name}, as {signer_title}",
            f"Email: {signer_email}",
        ])
        _add_body(doc, "Together referred to as the \"Parties\".")

    # ---- Art 1
    if is_fr:
        _add_section(doc, "Article 1 — Objet et contexte")
        _add_body(doc,
            "Le présent Accord définit les conditions dans lesquelles les Parties s'engagent à protéger "
            "la confidentialité des informations échangées dans le cadre des discussions précontractuelles relatives à :")
        _add_context_box(doc, mission_context)
        _add_body(doc,
            "Ces discussions peuvent inclure sans s'y limiter : échanges techniques, démonstrations, "
            "propositions commerciales, audits de systèmes existants, accès à des données d'exploitation "
            "ou à toute autre information sensible.")
    else:
        _add_section(doc, "Article 1 — Purpose and Context")
        _add_body(doc,
            "This Agreement sets out the conditions under which the Parties undertake to protect the confidentiality "
            "of information exchanged in connection with pre-contractual discussions regarding:")
        _add_context_box(doc, mission_context)
        _add_body(doc,
            "Such discussions may include, without limitation: technical exchanges, demonstrations, commercial proposals, "
            "audits of existing systems, access to operational data, and any other sensitive information.")

    # ---- Art 2
    if is_fr:
        _add_section(doc, "Article 2 — Définition des Informations Confidentielles")
        _add_body(doc,
            "Constituent des « Informations Confidentielles » toutes les informations, quelle qu'en soit la forme "
            "(orale, écrite, électronique, visuelle), communiquées par une Partie (la « Partie Divulgatrice ») "
            "à l'autre (la « Partie Réceptrice »), y compris notamment :")
        for item in [
            "les architectures techniques, codes sources, algorithmes propriétaires et configurations logicielles ;",
            "les modèles d'IA, prompts système, données d'entraînement, architectures d'agents et intégrations ;",
            "les données commerciales, financières, tarifaires, stratégiques et secrets d'affaires ;",
            "les processus métier, cahiers des charges, spécifications fonctionnelles et techniques ;",
            "les accès, identifiants, clés API, tokens et systèmes d'authentification ;",
            "les informations relatives aux fournisseurs, partenaires, clients et prospects ;",
            "toute information désignée comme confidentielle ou dont la nature confidentielle est raisonnablement évidente.",
        ]: _add_bullet(doc, item)
    else:
        _add_section(doc, "Article 2 — Definition of Confidential Information")
        _add_body(doc,
            "\"Confidential Information\" means all information, in any form (oral, written, electronic, visual), "
            "communicated by one Party (the \"Disclosing Party\") to the other (the \"Receiving Party\"), "
            "including but not limited to:")
        for item in [
            "technical architectures, source code, proprietary algorithms and software configurations;",
            "AI models, system prompts, training data, agent architectures and integrations;",
            "commercial, financial, pricing, strategic data and trade secrets;",
            "business processes, specifications, functional and technical requirements;",
            "credentials, API keys, tokens and authentication systems;",
            "information relating to suppliers, partners, customers and prospects;",
            "any information designated as confidential or whose confidential nature is reasonably apparent.",
        ]: _add_bullet(doc, item)

    # ---- Art 3
    if is_fr:
        _add_section(doc, "Article 3 — Obligations des Parties")
        _add_body(doc, "Chaque Partie Réceptrice s'engage à :")
        for item in [
            "traiter les Informations Confidentielles avec le même niveau de protection que ses propres informations confidentielles, et au minimum avec une diligence raisonnable ;",
            "ne pas divulguer les Informations Confidentielles à des tiers sans accord écrit préalable de la Partie Divulgatrice ;",
            "n'utiliser les Informations Confidentielles qu'aux seules fins des discussions précontractuelles visées à l'Article 1 ;",
            "limiter l'accès aux seules personnes ayant besoin d'en connaître, soumises à des obligations équivalentes ;",
            "informer sans délai la Partie Divulgatrice en cas de divulgation non autorisée, même accidentelle.",
        ]: _add_bullet(doc, item)
    else:
        _add_section(doc, "Article 3 — Obligations of the Parties")
        _add_body(doc, "Each Receiving Party undertakes to:")
        for item in [
            "treat Confidential Information with the same level of protection as its own confidential information, and at a minimum with reasonable care;",
            "not disclose Confidential Information to any third party without the prior written consent of the Disclosing Party;",
            "use Confidential Information solely for the purposes of the pre-contractual discussions referred to in Article 1;",
            "restrict access to persons who need to know it and who are bound by equivalent obligations;",
            "promptly notify the Disclosing Party of any unauthorised disclosure, even accidental.",
        ]: _add_bullet(doc, item)

    # ---- Art 4
    if is_fr:
        _add_section(doc, "Article 4 — Responsabilité pour les tiers et partenaires")
        _add_body(doc,
            "La Partie Réceptrice s'assure que ses partenaires, consultants, sous-traitants et filiales "
            "qui accèdent aux Informations Confidentielles sont liés par des obligations de confidentialité "
            "au moins aussi restrictives que celles du présent Accord.")
        _add_body_mixed(doc, [
            ("La Partie Réceptrice est strictement et solidairement responsable de tout manquement au présent Accord "
             "commis par ses partenaires, consultants ou filiales", True),
            (", comme si ce manquement était de son propre fait.", False),
        ])
    else:
        _add_section(doc, "Article 4 — Third-Party and Partner Liability")
        _add_body(doc,
            "The Receiving Party shall ensure that its partners, consultants, sub-contractors and affiliates "
            "who access Confidential Information are bound by confidentiality obligations at least as restrictive "
            "as those set out in this Agreement.")
        _add_body_mixed(doc, [
            ("The Receiving Party is strictly and jointly liable for any breach of this Agreement committed by its "
             "partners, consultants or affiliates", True),
            (", as if such breach were its own.", False),
        ])

    # ---- Art 5
    if is_fr:
        _add_section(doc, "Article 5 — Informations non confidentielles (Exclusions)")
        _add_body(doc, "Les obligations de l'Article 3 ne s'appliquent pas aux informations dont la Partie Réceptrice démontre :")
        for item in [
            "qu'elles étaient dans le domaine public au moment de leur communication, ou qu'elles y sont tombées sans faute de sa part ;",
            "qu'elles lui étaient déjà connues avant la communication par la Partie Divulgatrice ;",
            "qu'elles ont été développées indépendamment, sans recours aux Informations Confidentielles ;",
            "qu'elles ont été légalement obtenues d'un tiers sans obligation de confidentialité ;",
            "que leur divulgation est requise par une obligation légale ou judiciaire — dans ce cas, la Partie Réceptrice en informe préalablement la Partie Divulgatrice et limite la divulgation au strict nécessaire.",
        ]: _add_bullet(doc, item)
    else:
        _add_section(doc, "Article 5 — Exclusions")
        _add_body(doc, "The obligations in Article 3 shall not apply to information that the Receiving Party demonstrates:")
        for item in [
            "was in the public domain at the time of communication, or entered the public domain through no fault of the Receiving Party;",
            "was already known to the Receiving Party prior to communication by the Disclosing Party;",
            "was independently developed without use of the Confidential Information;",
            "was lawfully obtained from a third party free from any obligation of confidence;",
            "is required to be disclosed by law or court order — in which case the Receiving Party shall notify the Disclosing Party in advance and limit disclosure to the minimum necessary.",
        ]: _add_bullet(doc, item)

    # ---- Art 6
    if is_fr:
        _add_section(doc, "Article 6 — Durée")
        _add_body(doc,
            f"Le présent Accord entre en vigueur à sa signature et reste en vigueur pendant {nda_duration} "
            "à compter de la dernière divulgation d'Informations Confidentielles, sauf résiliation anticipée par accord écrit. "
            "Les obligations relatives aux informations constituant des secrets d'affaires au sens de la loi n° 2018-670 "
            "du 30 juillet 2018 survivent à l'expiration du présent Accord pour une durée indéfinie.")
    else:
        _add_section(doc, "Article 6 — Term")
        _add_body(doc,
            f"This Agreement takes effect upon signature and remains in force for {nda_duration} from the date of the last "
            "disclosure of Confidential Information, unless terminated earlier by written agreement. "
            "Obligations relating to information constituting trade secrets survive the expiry of this Agreement indefinitely.")

    # ---- Art 7 — Penalty box
    if is_fr:
        _add_section(doc, "Article 7 — Pénalités et dommages-intérêts")
        _add_body(doc,
            "Les Parties reconnaissent que toute violation du présent Accord causerait à la Partie Divulgatrice "
            "un préjudice irréparable pour lequel les seuls dommages-intérêts compensatoires pourraient être insuffisants.")
        _add_penalty_box(doc,
            "En cas de violation avérée du présent Accord, la Partie contrevenante s'engage à payer à la Partie lésée :",
            "250 000 € (deux cent cinquante mille euros) par violation constatée",
            "à titre de clause pénale au sens de l'article 1231-5 du Code civil, sans préjudice du droit de la "
            "Partie lésée à réclamer une compensation supplémentaire pour les pertes réelles excédant ce montant, "
            "ainsi que toute mesure d'injonction ou de référé.",
        )
        _add_body(doc, "Le paiement de cette pénalité ne dispense pas la Partie contrevenante de ses obligations au titre du présent Accord.")
    else:
        _add_section(doc, "Article 7 — Penalties and Damages")
        _add_body(doc,
            "The Parties acknowledge that any breach of this Agreement would cause the Disclosing Party irreparable harm "
            "for which monetary damages alone may be insufficient.")
        _add_penalty_box(doc,
            "In the event of a proven breach of this Agreement, the breaching Party shall pay the injured Party:",
            "€250,000 (two hundred and fifty thousand euros) per breach",
            "as liquidated damages, without prejudice to the injured Party's right to claim additional compensation "
            "for actual losses exceeding this amount, as well as any injunctive or other equitable relief.",
        )
        _add_body(doc, "Payment of this penalty does not release the breaching Party from its obligations under this Agreement.")

    # ---- Art 8
    if is_fr:
        _add_section(doc, "Article 8 — Non-sollicitation et non-contournement")
        _add_body(doc,
            "Pendant la durée du présent Accord et les vingt-quatre (24) mois suivant son terme, chaque Partie s'interdit "
            "de débaucher tout collaborateur ou sous-traitant de l'autre Partie ayant participé aux discussions, "
            "et de contacter directement les fournisseurs, partenaires ou prospects dont l'identité aurait été révélée "
            "dans le cadre de ces discussions. La violation entraîne une indemnité forfaitaire de 20 000 € par infraction, "
            "sans préjudice du recours à l'Article 7.")
    else:
        _add_section(doc, "Article 8 — Non-Solicitation and Non-Circumvention")
        _add_body(doc,
            "For the duration of this Agreement and twenty-four (24) months thereafter, each Party shall not solicit "
            "or hire any employee or sub-contractor of the other Party who participated in the discussions, "
            "nor contact directly any suppliers, partners or prospects whose identity was disclosed in the context of "
            "those discussions. A breach shall incur a fixed penalty of €20,000 per infringement, without prejudice to "
            "the application of Article 7.")

    # ---- Art 9
    if is_fr:
        _add_section(doc, "Article 9 — Propriété des Informations Confidentielles")
        _add_body(doc,
            "Le présent Accord ne confère à la Partie Réceptrice aucun droit de propriété intellectuelle, "
            "aucune licence, ni aucun autre droit sur les Informations Confidentielles. "
            "Toutes les Informations Confidentielles demeurent la propriété exclusive de la Partie Divulgatrice.")
    else:
        _add_section(doc, "Article 9 — Ownership of Confidential Information")
        _add_body(doc,
            "This Agreement does not grant the Receiving Party any intellectual property rights, licence, "
            "or other rights in respect of Confidential Information. All Confidential Information remains the "
            "exclusive property of the Disclosing Party.")

    # ---- Art 10
    if is_fr:
        _add_section(doc, "Article 10 — Retour et destruction des informations")
        _add_body(doc,
            "À la demande écrite de la Partie Divulgatrice, ou à l'expiration de l'Accord sans conclusion d'un contrat "
            "de prestation, la Partie Réceptrice restituera ou détruira tous supports contenant des Informations "
            "Confidentielles et certifiera par écrit cette exécution dans un délai de quinze (15) jours ouvrés.")
    else:
        _add_section(doc, "Article 10 — Return and Destruction of Information")
        _add_body(doc,
            "Upon written request by the Disclosing Party, or upon expiry of the Agreement without conclusion of a "
            "services contract, the Receiving Party shall return or destroy all media containing Confidential Information "
            "and certify such destruction in writing within fifteen (15) business days.")

    # ---- Art 11
    if is_fr:
        _add_section(doc, "Article 11 — Protection des données personnelles")
        _add_body(doc,
            "Si des données à caractère personnel sont communiquées dans le cadre des discussions, chaque Partie "
            "s'engage à les traiter conformément au RGPD (UE) 2016/679 et à la loi Informatique et Libertés modifiée, "
            "aux seules fins nécessaires aux discussions visées à l'Article 1.")
    else:
        _add_section(doc, "Article 11 — Data Protection")
        _add_body(doc,
            "Where personal data is communicated in the course of the discussions, each Party undertakes to process "
            "it in compliance with the GDPR (EU) 2016/679, solely for the purposes of the discussions referred to in Article 1.")

    # ---- Art 12
    if is_fr:
        _add_section(doc, "Article 12 — Droit applicable et juridiction compétente")
        _add_body(doc,
            "Le présent Accord est régi par le droit français. À défaut de résolution amiable dans un délai de "
            "trente (30) jours, tout litige sera soumis à la compétence exclusive du Tribunal des activités "
            "économiques de Paris.")
    else:
        _add_section(doc, "Article 12 — Governing Law and Jurisdiction")
        _add_body(doc,
            "This Agreement is governed by French law. In the absence of an amicable resolution within thirty (30) days, "
            "any dispute shall be submitted to the exclusive jurisdiction of the Commercial Court of Paris "
            "(Tribunal des activités économiques de Paris).")

    # ---- Art 13
    if is_fr:
        _add_section(doc, "Article 13 — Stipulations diverses")
        _add_body(doc,
            "Le présent Accord constitue l'intégralité de l'accord entre les Parties sur la confidentialité et "
            "ne constitue ni engagement à conclure un contrat de prestation, ni exclusivité de négociation. "
            "Toute modification fait l'objet d'un avenant écrit signé par les deux Parties. "
            "La nullité d'une stipulation n'affecte pas les autres.")
    else:
        _add_section(doc, "Article 13 — Miscellaneous")
        _add_body(doc,
            "This Agreement constitutes the entire agreement between the Parties with respect to confidentiality and "
            "does not constitute a commitment to enter into a services contract or an exclusivity of negotiation. "
            "Any amendment shall be made by a written addendum signed by both Parties. "
            "The invalidity of any provision shall not affect the remaining provisions.")

    # ---- Acceptance + Signatures
    if is_fr:
        _add_section(doc, "Acceptation valant signature")
        _add_body(doc,
            "Le Client déclare avoir pris connaissance et accepter sans réserve l'intégralité des stipulations "
            "du présent Accord, y compris la clause pénale de 250 000 € par violation (Article 7).")
        _add_signature_table(
            doc,
            "Pour le Client",
            "Pour Lucid-Lab",
            signer_name,
        )
    else:
        _add_section(doc, "Acceptance and Signatures")
        _add_body(doc,
            "By signing this Agreement, the Client confirms having read, understood and accepted all of its provisions, "
            "including the liquidated damages clause of €250,000 per breach (Article 7).")
        _add_signature_table(
            doc,
            "For the Client",
            "For Lucid-Lab",
            signer_name,
        )

    doc.save(str(out_path))
    tmp.unlink(missing_ok=True)
    return out_path


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_fr(**kwargs) -> Path:
    return _build_nda(lang="fr", **kwargs)

def build_en(**kwargs) -> Path:
    return _build_nda(lang="en", **kwargs)


# ---------------------------------------------------------------------------
# CLI — generate David Guzman NDA
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    today   = date.today()
    ref     = f"NDA-{today.strftime('%Y%m%d')}-ARTINCELL-00001"
    issued  = today.strftime("%d/%m/%Y")

    common = dict(
        doc_ref         = ref,
        issued_date     = issued,
        client_name     = "Artincell",
        client_address  = "artincell.com",
        client_reg      = "À compléter / To be completed",
        signer_name     = "David Guzman",
        signer_title    = "Directeur Général / CEO",
        signer_email    = "david@artincell.com",
        mission_context = (
            "Intégration de solutions d'intelligence artificielle pour l'expansion commerciale sur le marché européen — "
            "évaluation d'alternatives pour les fonctionnalités de suggested orders (XTE Chain), "
            "de détection de ventes perdues et de chat fournisseur.\n\n"
            "Integration of artificial intelligence solutions for European market expansion — "
            "evaluation of alternatives for suggested orders (XTE Chain), lost sales detection, "
            "and vendor chat features."
        ),
    )

    path_fr = build_fr(**common, nda_duration="5 ans",
                       out_path=OUT_DIR / f"{ref}-FR.docx")
    path_en = build_en(**common, nda_duration="5 years",
                       out_path=OUT_DIR / f"{ref}-EN.docx")

    print(f"FR: {path_fr}")
    print(f"EN: {path_en}")

