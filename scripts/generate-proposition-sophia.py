"""
Génère la Proposition d'accompagnement — BSP 37 / Sophia Kanouni
dans le même format que les propositions Clément et Fanny.

Run : python3 scripts/generate-proposition-sophia.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "client-presentations"
OUT.mkdir(parents=True, exist_ok=True)

FILENAME = "09_Proposition-Sophia-Kanouni-BSP37.docx"

# ── helpers ────────────────────────────────────────────────────────────────────

def _run(paragraph, text, bold=None, size_emu=None, color_hex=None, font_name=None):
    run = paragraph.add_run(text)
    if bold is not None:
        run.font.bold = bold
    if size_emu is not None:
        run.font.size = size_emu
    if color_hex is not None:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    if font_name is not None:
        run.font.name = font_name
    return run


def _set_spacing(paragraph, before_emu=None, after_emu=None, left_indent_emu=None):
    pf = paragraph.paragraph_format
    if before_emu is not None:
        pf.space_before = Emu(before_emu)
    if after_emu is not None:
        pf.space_after = Emu(after_emu)
    if left_indent_emu is not None:
        pf.left_indent = Emu(left_indent_emu)


def _shade_row(row, fill_hex):
    """Apply background colour to every cell in a row."""
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tcPr.append(shd)


def _set_table_borders(table):
    """Add thin borders to every cell."""
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CCCCCC")
        borders.append(el)
    tblPr.append(borders)


def _set_col_widths(table, widths_emu):
    for i, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = Emu(widths_emu[i])


def add_h1(doc, text):
    p = doc.add_paragraph()
    _run(p, text, bold=True, size_emu=190500, color_hex="000000", font_name="Syne")
    _set_spacing(p, before_emu=177800, after_emu=101600)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    _run(p, text, bold=True, size_emu=152400, color_hex="000000", font_name="Syne")
    _set_spacing(p, before_emu=127000, after_emu=76200)
    return p


def add_body(doc, text, space_after=50800):
    p = doc.add_paragraph()
    _run(p, text)
    _set_spacing(p, after_emu=space_after)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    _run(p, f"\u2022  {text}")
    _set_spacing(p, after_emu=25400, left_indent_emu=179705)
    return p


def add_calendar(doc, label, description):
    p = doc.add_paragraph()
    _run(p, label + "  ", bold=True)
    _run(p, description)
    _set_spacing(p, after_emu=38100)
    return p


def add_meta_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    _set_col_widths(table, [2_286_000, 5_943_600])
    _set_table_borders(table)
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        # alternating shade
        if i % 2 == 0:
            _shade_row(row, "F5F5F2")
        lp = row.cells[0].paragraphs[0]
        _run(lp, label, bold=True, size_emu=120650)
        vp = row.cells[1].paragraphs[0]
        _run(vp, value, size_emu=127000)
    _set_spacing(doc.add_paragraph(), after_emu=177800)  # spacer
    return table


def add_investment_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    _set_col_widths(table, [2_921_000, 5_308_600])
    _set_table_borders(table)
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        if i % 2 == 0:
            _shade_row(row, "F5F5F2")
        lp = row.cells[0].paragraphs[0]
        _run(lp, label, bold=True, size_emu=120650)
        vp = row.cells[1].paragraphs[0]
        _run(vp, value, size_emu=127000)
    return table


def add_signature_table(doc, client_name):
    table = doc.add_table(rows=4, cols=2)
    _set_col_widths(table, [4_114_800, 4_114_800])
    _set_table_borders(table)
    headers = [f"Pour {client_name}", "Pour Lucid-Lab"]
    labels = [
        ["Date :", "Date :"],
        ["Lieu :", "Lieu :"],
        [
            "Signature précédée de la mention\n« Lu et approuvé, bon pour accord »\u202f:",
            "Signature précédée de la mention\n« Lu et approuvé, bon pour accord »\u202f:",
        ],
    ]
    for j, h in enumerate(headers):
        p = table.rows[0].cells[j].paragraphs[0]
        _run(p, h, bold=True, size_emu=120650)
    for i, row_labels in enumerate(labels):
        row = table.rows[i + 1]
        for j, lbl in enumerate(row_labels):
            p = row.cells[j].paragraphs[0]
            _run(p, lbl, bold=True, size_emu=114300)
    return table


# ── document ───────────────────────────────────────────────────────────────────

LOGO = ROOT / "public" / "logo.png"

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Inches(0.9)
section.bottom_margin = Inches(0.9)
section.left_margin = Inches(1.1)
section.right_margin = Inches(1.1)

# Default font
style = doc.styles["Normal"]
style.font.name = "Inter"
style.font.size = Pt(10)

# ── Header (logo + title) ──────────────────────────────────────────────────────
hdr = section.header
hdr_para = hdr.paragraphs[0]
hdr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_logo = hdr_para.add_run()
run_logo.add_picture(str(LOGO), width=Emu(266700))  # ~0.29 inch / 7.4mm
run_txt = hdr_para.add_run("    Lucid-Lab    Proposition d'accompagnement — BSP 37")
run_txt.font.size = Pt(8.5)
run_txt.font.color.rgb = RGBColor.from_string("666666")

# ── Footer ─────────────────────────────────────────────────────────────────────
ftr = section.footer
ftr_para = ftr.paragraphs[0]
ftr_run = ftr_para.add_run(
    "Lucid-Lab · SAS au capital de 999 € · RCS Paris 104 672 050 · TVA FR 02 104 672 050 · info@lucid-lab.fr · lucid-lab.fr"
)
ftr_run.font.size = Pt(8)
ftr_run.font.color.rgb = RGBColor.from_string("999999")

# ── Title block ────────────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
_run(p_title, "PROPOSITION D'ACCOMPAGNEMENT", bold=True, size_emu=279400, color_hex="000000", font_name="Syne")
_set_spacing(p_title, before_emu=228600, after_emu=152400)

p_client = doc.add_paragraph()
_run(p_client, "BSP 37 — Sophia Kanouni", bold=True)
_set_spacing(p_client, after_emu=50800)

p_sub = doc.add_paragraph()
_run(p_sub, "Refonte + maintenance du site web bsp37.com et gestion Meta Ads à partir du 1er septembre 2026.", color_hex="666666")
_set_spacing(p_sub, after_emu=177800)

# ── Meta table ─────────────────────────────────────────────────────────────────
add_meta_table(doc, [
    ("Client",     "BSP 37 — Sophia Kanouni"),
    ("Date",       "23 mai 2026"),
    ("Contact",    "Lucid-Lab — info@lucid-lab.fr"),
    ("Référence",  "PROP-2026-009"),
    ("Validité",   "30 jours à compter de la date d'émission"),
])

# ── Objectif ───────────────────────────────────────────────────────────────────
add_h1(doc, "Objectif")
add_body(doc,
    "BSP 37 est un groupement familial actif dans l'artisanat du bâtiment depuis 1989. "
    "L'objectif de cet accompagnement est double : moderniser la présence digitale de l'entreprise "
    "avec une refonte complète du site web bsp37.com, puis ouvrir un canal d'acquisition pour "
    "l'activité film solaire via Meta Ads à partir du 1er septembre 2026.",
    space_after=177800,
)

# ── Périmètre ──────────────────────────────────────────────────────────────────
add_h1(doc, "Périmètre")

add_h2(doc, "Refonte du site web — bsp37.com (incluse dans le forfait mensuel)")
for item in [
    "Audit du site Wix existant et recommandations.",
    "Refonte complète sur technologie moderne : meilleure performance, SEO et contrôle total sans abonnement Wix.",
    "Intégration du catalogue produits, galerie de réalisations, pages de services et formulaires de contact.",
    "Déploiement et configuration du nom de domaine.",
    "Maintenance mensuelle : mises à jour, corrections et suivi des performances SEO.",
]:
    add_bullet(doc, item)

add_h2(doc, "Meta Ads — Film solaire (à partir du 1er septembre 2026)")
for item in [
    "Création et configuration du compte Meta Ads Facebook / Instagram.",
    "Définition de la cible locale : zone géographique, audiences et centres d'intérêt.",
    "Création de deux publicités Meta Ads pour le lancement : 250 € HT en frais unique.",
    "Gestion mensuelle des campagnes : optimisation, tests visuels, reporting mensuel.",
    "Mise en place d'un suivi des leads : formulaire dédié, traçabilité et tableau de bord partagé.",
    "Budget publicitaire Meta non inclus, réglé directement par Sophia auprès de Meta.",
]:
    add_bullet(doc, item)

# ── Calendrier ─────────────────────────────────────────────────────────────────
add_h1(doc, "Calendrier")
add_calendar(doc, "Semaine 1", "Audit du site Wix existant, cadrage du nouveau site et préparation du brief Meta Ads.")
add_calendar(doc, "Semaines 2–3", "Refonte du site web : structure, design, contenus et intégrations.")
add_calendar(doc, "Semaine 4", "Mise en ligne du site et validation du parcours de génération de leads film solaire.")
add_calendar(doc, "À partir du 1er septembre 2026", "Création des deux publicités, lancement et gestion mensuelle des campagnes Meta Ads.")

# ── Investissement ─────────────────────────────────────────────────────────────
add_h1(doc, "Investissement")
add_investment_table(doc, [
    ("Forfait mensuel — Site web",    "150 € HT / mois"),
    ("Gestion Meta Ads",              "100 € HT / mois à partir du 1er septembre 2026"),
    ("Forfait mensuel total",         "250 € HT / mois à partir du 1er septembre 2026 (site web + gestion Meta Ads)"),
    ("Création de 2 publicités",      "250 € HT en frais unique"),
    ("TVA (20 %)",                    "Applicable en sus des prix HT ci-dessus"),
    ("Conditions de paiement",        "Virement SEPA exclusivement, à l'avance — voir CGV article 5"),
    ("Engagement (forfait mensuel)",  "Ferme et irrévocable sur 12 mois — clause pénale en cas de rupture anticipée"),
    ("Pénalités de retard",           "Intérêts BCE +10 pts (≈14,5 %/an) + 40 € (D. 441-5 Cciv) + 250 € par facture"),
])

# ── Coûts exclus ───────────────────────────────────────────────────────────────
_set_spacing(doc.add_paragraph(), after_emu=101600)  # spacer after table
add_h1(doc, "Coûts exclus du forfait")
add_body(doc,
    "Le prix de la Prestation rémunère exclusivement le travail de Lucid-Lab. "
    "Restent à votre charge directe les coûts variables suivants, par nature dépendants de votre activité et de votre usage :",
    space_after=50800,
)
for item in [
    "Budget publicitaire Meta Ads, réglé directement par Sophia auprès de Meta ;",
    "Abonnements aux outils tiers utilisés pour votre compte (hébergeur, nom de domaine, etc.) ;",
    "Frais de déplacement éventuels, sur accord préalable écrit.",
]:
    add_bullet(doc, item)

add_body(doc,
    "Une estimation indicative de ces coûts vous est fournie sur demande, sans engagement de plafond.",
    space_after=50800,
)
add_body(doc,
    "Les présentes conditions financières sont régies par les Conditions Générales de Vente de Lucid-Lab "
    "— Version 1.0 accessibles à https://lucid-lab.fr/cgv.",
    space_after=177800,
)

# ── Prochaines étapes ──────────────────────────────────────────────────────────
add_h1(doc, "Prochaines étapes")
for item in [
    "Confirmer la date de démarrage du site et le lancement Meta Ads au 1er septembre 2026.",
    "Fournir les accès au site Wix, au domaine bsp37.com et au compte Meta Business.",
    "Valider le budget publicitaire Meta Ads mensuel, distinct des frais de gestion Lucid-Lab.",
    "Signer le Bon de Commande et procéder au premier paiement pour démarrage.",
]:
    add_bullet(doc, item)

# ── Cadre contractuel ─────────────────────────────────────────────────────────
add_h1(doc, "Cadre contractuel et démarrage")
add_body(doc,
    "Démarrage des Prestations : après signature du Bon de Commande accompagnant la présente Proposition "
    "et réception effective du premier paiement sur le compte bancaire de Lucid-Lab "
    "(IBAN FR76 1732 8844 0043 2662 8862 178 — BIC SWNBFR22).",
    space_after=50800,
)
add_body(doc,
    "Documents contractuels remis : (i) la présente Proposition, (ii) le Bon de Commande à signer, "
    "(iii) le Contrat de Prestation, (iv) les CGV v1.0, (v) le cas échéant l'Accord de sous-traitance des données (DPA).",
    space_after=50800,
)
add_body(doc,
    "Couverture assurantielle : Lucid-Lab est assurée en Responsabilité Civile Professionnelle auprès de Hiscox SA "
    "(police RCPLP 3695175450, plafond 100 000 € par période d'assurance, monde entier hors USA/Canada).",
    space_after=177800,
)

# ── Acceptation ────────────────────────────────────────────────────────────────
add_h1(doc, "Acceptation valant commande ferme")
add_body(doc, "Le Client déclare avoir pris connaissance et accepter sans réserve :", space_after=50800)
for item in [
    "le périmètre, les livrables et le calendrier décrits ci-dessus ;",
    "la modalité tarifaire retenue et l'engagement de durée associé ;",
    "les Conditions Générales de Vente Lucid-Lab v1.0 accessibles sur lucid-lab.fr/cgv ;",
    "l'obligation de paiement préalable au démarrage des Prestations (art. 5 CGV).",
]:
    add_bullet(doc, item)

add_body(doc,
    "La signature de la présente Proposition vaut acceptation de l'offre, formation du Contrat et commande ferme, "
    "et déclenche l'émission d'une facture pro forma puis l'exécution des Prestations à compter de la réception effective du premier paiement.",
    space_after=101600,
)

p_modal = doc.add_paragraph()
_run(p_modal, "Modalité retenue : ☑ Forfait mensuel site + gestion Meta Ads à 250 € HT/mois à partir du 1er septembre 2026 + frais unique de 250 € HT pour deux publicités Meta Ads.", bold=True)
_set_spacing(p_modal, after_emu=76200)

p_mention = doc.add_paragraph()
_run(p_mention, "Mention manuscrite obligatoire : « Lu et approuvé — Bon pour accord et commande ferme — "
    "Prix : [Total TTC] € — Modalité : [forfait mensuel site + gestion Meta Ads + frais unique deux publicités] ».", bold=True)
_set_spacing(p_mention, after_emu=177800)

# ── Signature ──────────────────────────────────────────────────────────────────
add_signature_table(doc, "BSP 37 / Sophia Kanouni")

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = OUT / FILENAME
doc.save(str(out_path))
print(f"✓ Sauvegardé : {out_path}")
