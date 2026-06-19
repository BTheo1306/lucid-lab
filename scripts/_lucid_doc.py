"""
Shared helpers for Lucid-Lab branded .docx documents.

Design goals:
- A4 page size, fixed margins, predictable pagination.
- One single logo placement (cover page only), aligned via table.
- Clean rubrics, minimal decoration, straight-to-the-point typography.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor, Inches, Emu


ROOT = Path(__file__).resolve().parents[1]
LOGO_PATH = ROOT / "public" / "logo.png"

INK = "0A0A0A"
MUTED = "555555"
SOFT = "DDDDDD"
ACCENT = "FFB451"  # amber, used as a thin rule only


def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.strip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_font(run, *, family="Inter", size=None, bold=None, italic=None, color=None):
    run.font.name = family
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), family)
    rFonts.set(qn("w:hAnsi"), family)
    rFonts.set(qn("w:cs"), family)
    rFonts.set(qn("w:eastAsia"), family)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = _rgb(color)


def _keep_with_next(paragraph, value=True):
    pPr = paragraph._p.get_or_add_pPr()
    el = pPr.find(qn("w:keepNext"))
    if el is None:
        el = OxmlElement("w:keepNext")
        pPr.append(el)
    el.set(qn("w:val"), "1" if value else "0")


def _keep_together(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    el = pPr.find(qn("w:keepLines"))
    if el is None:
        el = OxmlElement("w:keepLines")
        pPr.append(el)


def _clear(paragraph):
    paragraph._element.clear_content()


def _configure_page(section, first_page_header=True):
    # A4 portrait
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)
    if first_page_header:
        # different first-page header so the cover has no header logo/text
        sectPr = section._sectPr
        titlePg = sectPr.find(qn("w:titlePg"))
        if titlePg is None:
            titlePg = OxmlElement("w:titlePg")
            sectPr.append(titlePg)


def _add_header_footer(section):
    # Standard header (not first page): wordmark only, no logo image.
    header = section.header
    p = header.paragraphs[0]
    _clear(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Lucid-Lab")
    _set_font(run, family="Inter", size=9, bold=True, color=INK)
    p.paragraph_format.space_after = Pt(0)

    # First-page header: empty (clean cover).
    first_header = section.first_page_header
    fp = first_header.paragraphs[0]
    _clear(fp)
    fp.add_run("")

    # Footer with page number + contact.
    footer = section.footer
    fp_p = footer.paragraphs[0]
    _clear(fp_p)
    fp_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left = fp_p.add_run("lucid-lab.fr · info@lucid-lab.fr")
    _set_font(left, family="Inter", size=8, color="999999")
    tab = fp_p.add_run("\t")
    _set_font(tab, family="Inter", size=8, color="999999")
    pPr = fp_p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_el = OxmlElement("w:tab")
    tab_el.set(qn("w:val"), "right")
    tab_el.set(qn("w:pos"), "9000")
    tabs.append(tab_el)
    pPr.append(tabs)
    # PAGE field
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run = fp_p.add_run()
    _set_font(page_run, family="Inter", size=8, color="999999")
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_end)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Inter"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Inter")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _rgb(INK)


def init_document() -> _Document:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    _configure_page(section)
    _add_header_footer(section)
    return doc


# ---------- Building blocks ----------

def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_cover(doc, *, eyebrow: str, title: str, subtitle: str, meta: list[tuple[str, str]]):
    # Large logo + wordmark, ONE single placement, aligned in a table
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    
    # Remove borders
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    c0 = table.cell(0, 0)
    c1 = table.cell(0, 1)
    c0.width = Inches(0.55)
    c1.width = Inches(5.0)
    c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p0 = c0.paragraphs[0]
    _clear(p0)
    p0.paragraph_format.space_after = Pt(0)
    p0.paragraph_format.space_before = Pt(0)
    logo_run = p0.add_run()
    from docx.shared import Cm
    logo_run.add_picture(str(LOGO_PATH), width=Inches(0.40))

    p1 = c1.paragraphs[0]
    _clear(p1)
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.space_before = Pt(0)
    # The padding is effectively on the left due to the column width, no extra spaces needed
    name_run = p1.add_run("Lucid-Lab")
    _set_font(name_run, family="Syne", size=17, bold=True, color=INK)

    # Empty paragraph for spacing
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(28)

    if eyebrow:
        eb = doc.add_paragraph()
        eb.paragraph_format.space_after = Pt(8)
        run = eb.add_run(eyebrow.upper())
        _set_font(run, family="Inter", size=8.5, bold=True, color=MUTED)

    # Title
    t = doc.add_paragraph()
    t.paragraph_format.space_after = Pt(10)
    t.paragraph_format.line_spacing = 1.05
    run = t.add_run(title)
    _set_font(run, family="Inter", size=32, bold=True, color=INK)
    _keep_with_next(t)

    if subtitle:
        s = doc.add_paragraph()
        s.paragraph_format.space_after = Pt(28)
        s.paragraph_format.line_spacing = 1.45
        run = s.add_run(subtitle)
        _set_font(run, family="Inter", size=12, color=MUTED)

    # Thin accent rule (single subtle accent, not a band)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    pPr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT)
    pbdr.append(bottom)
    pPr.append(pbdr)

    # Meta as simple key/value lines (no table)
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        lab = p.add_run(f"{label.upper()}  ")
        _set_font(lab, family="Inter", size=8, bold=True, color=MUTED)
        val = p.add_run(value)
        _set_font(val, family="Inter", size=10.5, color=INK)


def add_section(doc, title: str, *, new_page=False):
    if new_page:
        page_break(doc)
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(10)
    h.paragraph_format.line_spacing = 1.1
    run = h.add_run(title)
    _set_font(run, family="Inter", size=18, bold=True, color=INK)
    _keep_with_next(h)
    return h


def add_paragraph(doc, text: str, *, bold=False, color=INK, after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.45
    run = p.add_run(text)
    _set_font(run, family="Inter", size=10.5, bold=bold, color=color)
    return p


def add_subheading(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_font(run, family="Inter", size=11, bold=True, color=INK)
    _keep_with_next(p)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.35
    _clear(p)
    run = p.add_run(text)
    _set_font(run, family="Inter", size=10.5, color=INK)
    _keep_together(p)
    return p


def add_kv_row(doc, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    lab = p.add_run(f"{label}  ")
    _set_font(lab, family="Inter", size=10.5, bold=True, color=INK)
    val = p.add_run(value)
    _set_font(val, family="Inter", size=10.5, color=MUTED)
    return p

def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set cell borders. Each param: (color_hex, size_in_8pt_units) or None to clear."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side_name, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        existing = tcBorders.find(qn(f"w:{side_name}"))
        if existing is not None:
            tcBorders.remove(existing)
        el = OxmlElement(f"w:{side_name}")
        if val is None:
            el.set(qn("w:val"), "none")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(val[1]))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val[0])
        tcBorders.append(el)


def add_timeline(doc, milestones: list[dict], usable_width_emu: int = 5_548_320):
    """
    Draws a visual horizontal timeline suitable for project roadmaps.

    milestones: list of dict with:
        date: str           — bold label above the dot
        label: str|None     — optional tag below date in orange (e.g. 'DÉMARRAGE')
        lines: list[str]    — short description lines below the dot

    First and last milestones get an orange dot; intermediate ones are dark.
    """
    n = len(milestones)
    col_w = usable_width_emu // n

    table = doc.add_table(rows=3, cols=n)
    table.autofit = False

    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tbl_borders.append(el)
    tblPr.append(tbl_borders)

    for j, ms in enumerate(milestones):
        is_accent = (j == 0 or j == n - 1)

        # Row 0 : date + optional tag
        c0 = table.cell(0, j)
        c0.width = Emu(col_w)
        _set_cell_border(c0, top=None, bottom=None, left=None, right=None)
        p = c0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(ms["date"])
        _set_font(r, family="Inter", size=9, bold=True, color=INK)
        if ms.get("label"):
            pl = c0.add_paragraph()
            pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pl.paragraph_format.space_before = Pt(0)
            pl.paragraph_format.space_after = Pt(2)
            rl = pl.add_run(ms["label"].upper())
            _set_font(rl, family="Inter", size=7, bold=True, color=ACCENT)

        # Row 1 : dot; top border = the connecting horizontal line
        c1 = table.cell(1, j)
        c1.width = Emu(col_w)
        _set_cell_border(c1, top=(SOFT, 6), bottom=None, left=None, right=None)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run("●")
        _set_font(r1, family="Inter", size=9, color=ACCENT if is_accent else INK)

        # Row 2 : description lines
        c2 = table.cell(2, j)
        c2.width = Emu(col_w)
        _set_cell_border(c2, top=None, bottom=None, left=None, right=None)
        for k, line in enumerate(ms.get("lines", [])):
            p2 = c2.paragraphs[0] if k == 0 else c2.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(2)
            r2 = p2.add_run(line)
            _set_font(r2, family="Inter", size=8, color=MUTED)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(18)


def add_signature_block(doc, client_name: str):
    h = add_section(doc, "Validation & Signatures", new_page=False)
    h.paragraph_format.space_before = Pt(30)
    
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    
    # Remove borders
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    c0 = table.cell(0, 0)
    c1 = table.cell(0, 1)
    c0.width = Inches(3.2)
    c1.width = Inches(3.2)
    
    for cell, entity in [(c0, f"Pour {client_name}"), (c1, "Pour Lucid-Lab")]:
        p = cell.paragraphs[0]
        _clear(p)
        r = p.add_run(entity)
        _set_font(r, family="Inter", size=10.5, bold=True, color=INK)
        
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(8)
        p2.paragraph_format.space_after = Pt(36)
        r2 = p2.add_run("Date : \n\nSignature :")
        _set_font(r2, family="Inter", size=10.5, color=MUTED)
        
        # Make sure signature block stays together
        _keep_together(p)
        _keep_together(p2)
