"""
Génère la Proposition d'accompagnement — Sinibaldi Agency / Permis de Construire IA
dans le même format que les propositions Sophia, Fanny, Clément existantes.

Run : python3 scripts/generate-proposition-sinibaldi-pc.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "client-presentations"
OUT.mkdir(parents=True, exist_ok=True)

FILENAME = "10_Proposition-Sinibaldi-Agency-PC-IA.docx"

LOGO = ROOT / "lucid-lab-brand" / "01-logo" / "png" / "logo-1024x1024-black-on-white.png"


# ── helpers ────────────────────────────────────────────────────────────────────

def _run(paragraph, text, bold=None, size_emu=None, color_hex=None, font_name=None):
    run = paragraph.add_run(text)
    if bold is not None:
        run.font.bold = bold
    if size_emu is not None:
        run.font.size = size_emu
    if color_hex is not None:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    if font_name is not None:
        run.font.name = font_name
    return run


def _set_spacing(paragraph, before_emu=None, after_emu=None, left_indent_emu=None):
    pf = paragraph.paragraph_format
    if before_emu is not None:
        pf.space_before = Emu(before_emu)
    if after_emu is not None:
        pf.space_after = Emu(after_emu)
    if left_indent_emu is not None:
        pf.left_indent = Emu(left_indent_emu)


def _shade_row(row, fill_hex):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tcPr.append(shd)


def _set_table_borders(table):
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CCCCCC")
        borders.append(el)
    tblPr.append(borders)


def _set_col_widths(table, widths_emu):
    for i, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = Emu(widths_emu[i])


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set cell borders. Each param: (color_hex, size_in_8pt_units) or None to clear."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side_name, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        existing = tcBorders.find(qn(f"w:{side_name}"))
        if existing is not None:
            tcBorders.remove(existing)
        el = OxmlElement(f"w:{side_name}")
        if val is None:
            el.set(qn("w:val"), "none")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(val[1]))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val[0])
        tcBorders.append(el)


def add_timeline(doc, milestones):
    """
    Draws a visual horizontal timeline.
    milestones: list of dict with:
        date: str           — bold label above the dot
        label: str|None     — optional tag below date in orange (e.g. 'DÉMARRAGE')
        lines: list[str]    — short description lines below the dot
    """
    USABLE_W = 5_548_320  # EMU — A4 with 1.1" margins
    n = len(milestones)
    col_w = USABLE_W // n

    table = doc.add_table(rows=3, cols=n)
    table.autofit = False

    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tbl_borders.append(el)
    tblPr.append(tbl_borders)

    for j, ms in enumerate(milestones):
        is_accent = (j == 0 or j == n - 1)

        # Row 0 : date + optional tag
        c0 = table.cell(0, j)
        c0.width = Emu(col_w)
        _set_cell_border(c0, top=None, bottom=None, left=None, right=None)
        p = c0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Emu(0)
        p.paragraph_format.space_after = Emu(19050)
        r = p.add_run(ms["date"])
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string("0A0A0A")
        r.font.name = "Inter"
        if ms.get("label"):
            pl = c0.add_paragraph()
            pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pl.paragraph_format.space_before = Emu(0)
            pl.paragraph_format.space_after = Emu(12700)
            rl = pl.add_run(ms["label"].upper())
            rl.bold = True
            rl.font.size = Pt(7)
            rl.font.color.rgb = RGBColor.from_string("FFB451")
            rl.font.name = "Inter"

        # Row 1 : dot on horizontal line (top border = the line)
        c1 = table.cell(1, j)
        c1.width = Emu(col_w)
        _set_cell_border(c1, top=("D0D0D0", 6), bottom=None, left=None, right=None)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Emu(12700)
        p1.paragraph_format.space_after = Emu(12700)
        r1 = p1.add_run("●")
        r1.font.size = Pt(9)
        r1.font.name = "Inter"
        r1.font.color.rgb = RGBColor.from_string("FFB451" if is_accent else "0A0A0A")

        # Row 2 : description lines
        c2 = table.cell(2, j)
        c2.width = Emu(col_w)
        _set_cell_border(c2, top=None, bottom=None, left=None, right=None)
        for k, line in enumerate(ms.get("lines", [])):
            p2 = c2.paragraphs[0] if k == 0 else c2.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Emu(0)
            p2.paragraph_format.space_after = Emu(12700)
            r2 = p2.add_run(line)
            r2.font.size = Pt(8)
            r2.font.color.rgb = RGBColor.from_string("555555")
            r2.font.name = "Inter"

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Emu(127000)


def add_h1(doc, text):
    p = doc.add_paragraph()
    _run(p, text, bold=True, size_emu=190500, color_hex="000000", font_name="Syne")
    _set_spacing(p, before_emu=177800, after_emu=101600)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    _run(p, text, bold=True, size_emu=152400, color_hex="000000", font_name="Syne")
    _set_spacing(p, before_emu=127000, after_emu=76200)
    return p


def add_body(doc, text, space_after=50800):
    p = doc.add_paragraph()
    _run(p, text)
    _set_spacing(p, after_emu=space_after)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    _run(p, f"•  {text}")
    _set_spacing(p, after_emu=25400, left_indent_emu=179705)
    return p


def add_calendar(doc, label, description):
    p = doc.add_paragraph()
    _run(p, label + "  ", bold=True)
    _run(p, description)
    _set_spacing(p, after_emu=38100)
    return p


def add_meta_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    _set_col_widths(table, [2_286_000, 5_943_600])
    _set_table_borders(table)
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        if i % 2 == 0:
            _shade_row(row, "F5F5F2")
        lp = row.cells[0].paragraphs[0]
        _run(lp, label, bold=True, size_emu=120650)
        vp = row.cells[1].paragraphs[0]
        _run(vp, value, size_emu=127000)
    _set_spacing(doc.add_paragraph(), after_emu=177800)
    return table


def add_investment_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    _set_col_widths(table, [2_921_000, 5_308_600])
    _set_table_borders(table)
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        if i % 2 == 0:
            _shade_row(row, "F5F5F2")
        lp = row.cells[0].paragraphs[0]
        _run(lp, label, bold=True, size_emu=120650)
        vp = row.cells[1].paragraphs[0]
        _run(vp, value, size_emu=127000)
    return table


def add_signature_table(doc, client_name):
    table = doc.add_table(rows=4, cols=2)
    _set_col_widths(table, [4_114_800, 4_114_800])
    _set_table_borders(table)
    headers = [f"Pour {client_name}", "Pour Lucid-Lab"]
    labels = [
        ["Date :", "Date :"],
        ["Lieu :", "Lieu :"],
        [
            "Signature précédée de la mention\n« Lu et approuvé, bon pour accord » :",
            "Signature précédée de la mention\n« Lu et approuvé, bon pour accord » :",
        ],
    ]
    for j, h in enumerate(headers):
        p = table.rows[0].cells[j].paragraphs[0]
        _run(p, h, bold=True, size_emu=120650)
    for i, row_labels in enumerate(labels):
        row = table.rows[i + 1]
        for j, lbl in enumerate(row_labels):
            p = row.cells[j].paragraphs[0]
            _run(p, lbl, bold=True, size_emu=114300)
    return table


# ── document ───────────────────────────────────────────────────────────────────

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(0.9)
section.bottom_margin = Inches(0.9)
section.left_margin = Inches(1.1)
section.right_margin = Inches(1.1)

style = doc.styles["Normal"]
style.font.name = "Inter"
style.font.size = Pt(10)

# ── Header ─────────────────────────────────────────────────────────────────────
hdr = section.header
hdr_para = hdr.paragraphs[0]
hdr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_logo = hdr_para.add_run()
run_logo.add_picture(str(LOGO), width=Emu(266700))
run_txt = hdr_para.add_run("    Lucid-Lab    Proposition d'accompagnement · Sinibaldi Agency")
run_txt.font.size = Pt(8.5)
run_txt.font.color.rgb = RGBColor.from_string("666666")

# ── Footer ─────────────────────────────────────────────────────────────────────
ftr = section.footer
ftr_para = ftr.paragraphs[0]
ftr_run = ftr_para.add_run(
    "Lucid-Lab · SAS au capital de 999 € · RCS Paris 104 672 050 · TVA FR 02 104 672 050 · info@lucid-lab.fr · lucid-lab.fr"
)
ftr_run.font.size = Pt(8)
ftr_run.font.color.rgb = RGBColor.from_string("999999")

# ── Title block ────────────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
_run(p_title, "PROPOSITION D'ACCOMPAGNEMENT", bold=True, size_emu=279400, color_hex="000000", font_name="Syne")
_set_spacing(p_title, before_emu=228600, after_emu=152400)

p_client = doc.add_paragraph()
_run(p_client, "Sinibaldi Agency, Clément Sinibaldi", bold=True)
_set_spacing(p_client, after_emu=50800)

p_sub = doc.add_paragraph()
_run(p_sub, "Développement d'une plateforme IA de génération automatique de dossiers de permis de construire. Phase 1 B2B.", color_hex="666666")
_set_spacing(p_sub, after_emu=177800)

# ── Meta table ─────────────────────────────────────────────────────────────────
add_meta_table(doc, [
    ("Client",    "Sinibaldi Agency, Clément Sinibaldi"),
    ("Date",      "9 juin 2026"),
    ("Contact",   "Lucid-Lab, info@lucid-lab.fr"),
    ("Référence", "PROP-2026-010"),
    ("Validité",  "30 jours à compter de la date d'émission"),
])

# ── Objectif ───────────────────────────────────────────────────────────────────
add_h1(doc, "Objectif")
add_body(doc,
    "Sinibaldi Agency produit chaque mois des dossiers complets de Déclarations Préalables et "
    "Permis de Construire pour ses clients : plans de situation, plans cadastraux, notices architecturales, "
    "élévations de façades (état existant et état projet), rendus. Chaque dossier mobilise plusieurs "
    "heures de production manuelle.",
)
add_body(doc,
    "L'objectif de cet accompagnement est de concevoir et de développer une plateforme IA capable de "
    "générer automatiquement l'ensemble de ces pièces graphiques et écrites, à partir des données de "
    "projet : adresse, références cadastrales, description des travaux, photos de l'existant. L'IA "
    "produit les plans, les élévations, les rendus et les notices en quelques minutes, au standard "
    "attendu par les services instructeurs.",
)
add_body(doc,
    "La stratégie retenue est une approche en deux phases : Phase 1 B2B (3 mois, 9 000 € HT) "
    "pour développer et valider le MVP avec Sinibaldi Agency et un premier panel professionnel. "
    "Phase 2 B2C, budget à définir conjointement à l'issue de la Phase 1, pour un lancement "
    "auprès des particuliers souhaitant gérer eux-mêmes leur dossier de travaux.",
    space_after=177800,
)

# ── Périmètre ──────────────────────────────────────────────────────────────────
add_h1(doc, "Périmètre")

add_h2(doc, "Moteur IA : génération des pièces du dossier")
for item in [
    "Plans de situation et plans cadastraux générés automatiquement à partir de l'adresse et des références parcellaires (intégration des données cadastrales officielles : Géoportail, API cadastre).",
    "Notices architecturales rédigées automatiquement à partir de la description du projet : nature des travaux, matériaux, dimensions, caractéristiques de l'existant.",
    "Élévations de façades (état existant et état projet) générées par IA à partir de photos de la construction et de la description des modifications.",
    "Rendus architecturaux produits par IA pour illustrer le dossier.",
    "Pré-remplissage des formulaires CERFA (13703, 16702, 13406…) à partir des données saisies.",
]:
    add_bullet(doc, item)

add_h2(doc, "Interface et workflow")
for item in [
    "Interface de saisie structurée : adresse, références cadastrales, type de travaux, photos de l'existant, description du projet.",
    "Génération du dossier complet en PDF au format attendu par les services instructeurs.",
    "Relecture et ajustement possibles par l'architecte avant envoi.",
    "Archivage et gestion des dossiers par client / projet.",
]:
    add_bullet(doc, item)

add_h2(doc, "Go-to-market B2B et préparation Phase 2")
for item in [
    "Onboarding de Sinibaldi Agency comme premier client pilote et co-développeur produit.",
    "Identification et approche d'un panel complémentaire : cabinets d'architecture d'intérieur, agences immobilières avec activité travaux, entreprises générales de rénovation.",
    "Collecte structurée des retours terrain pour prioriser les itérations.",
    "Définition du modèle tarifaire B2B (abonnement par siège ou à la consommation par dossier) et recommandations pour le périmètre de la Phase 2 B2C.",
]:
    add_bullet(doc, item)

# ── Calendrier ─────────────────────────────────────────────────────────────────
add_h1(doc, "Calendrier")
add_timeline(doc, [
    {"date": "Juillet",    "label": "DÉMARRAGE", "lines": ["Cadrage + socle"]},
    {"date": "Mi-juillet", "label": None,         "lines": ["Plans cadastraux"]},
    {"date": "Août",       "label": None,         "lines": ["Élévations + rendus"]},
    {"date": "Mi-août",    "label": None,         "lines": ["Interface + PDF"]},
    {"date": "Septembre",  "label": None,         "lines": ["Tests + onboarding"]},
    {"date": "Fin sept.",  "label": "LIVRAISON",  "lines": ["Rapport + reco. B2C"]},
])
add_calendar(doc, "Mois 1",
    "Architecture de la plateforme, intégration des API cadastrales, développement du moteur de génération "
    "des plans de situation et cadastraux, premiers prototypes d'élévations.")
add_calendar(doc, "Mois 2",
    "Notices architecturales automatisées, élévations façades (existant + projet), rendus, pré-remplissage CERFA, "
    "interface de saisie et génération PDF. Test interne sur dossiers réels Sinibaldi Agency.")
add_calendar(doc, "Mois 3",
    "Ajustements sur retours terrain, onboarding des premiers clients B2B externes, collecte des métriques, "
    "rapport de phase et recommandations Phase 2 B2C.")

# ── Investissement ─────────────────────────────────────────────────────────────
add_h1(doc, "Investissement")
add_investment_table(doc, [
    ("Forfait Phase 1 B2B (3 mois)",     "9 000,00 € HT"),
    ("TVA (20 %)",                       "1 800,00 €"),
    ("Total TTC",                        "10 800,00 € TTC"),
    ("Modalité retenue",                 "☑ Forfait fixe Phase 1 B2B, 3 mois"),
    ("Échéancier",                       "30 % à la signature (3 240,00 € TTC) · 40 % à mi-parcours M1 (4 320,00 € TTC) · 30 % à la livraison M3 (3 240,00 € TTC)"),
    ("Phase 2 B2C",                       "Budget à définir conjointement à l'issue de la Phase 1, selon résultats et orientations validées"),
    ("Conditions de paiement",           "Virement SEPA exclusivement, à l'avance (voir CGV article 5)"),
    ("Pénalités de retard",              "Intérêts BCE +10 pts (≈14,5 %/an) + 40 € (D. 441-5 Cciv) + 250 € par facture"),
])

# ── Coûts exclus ───────────────────────────────────────────────────────────────
_set_spacing(doc.add_paragraph(), after_emu=101600)
add_h1(doc, "Coûts exclus du forfait")
add_body(doc,
    "Le prix de la Prestation rémunère exclusivement le travail de Lucid-Lab. "
    "Restent à votre charge directe les coûts variables suivants, par nature dépendants de l'usage réel :",
    space_after=50800,
)
for item in [
    "Crédits API des modèles IA de génération d'images et de texte (OpenAI, Mistral, Stability, ou équivalents) ;",
    "Coûts d'hébergement de la plateforme selon usage réel (serveur, stockage, bande passante) ;",
    "Accès aux API cadastrales et géographiques au-delà des quotas gratuits ;",
    "Budget marketing ou acquisition B2B éventuel ;",
    "Frais de déplacement éventuels, sur accord préalable écrit.",
]:
    add_bullet(doc, item)

add_body(doc,
    "Une estimation indicative de ces coûts vous est fournie sur demande, sans engagement de plafond.",
    space_after=50800,
)
add_body(doc,
    "Les présentes conditions financières sont régies par les Conditions Générales de Vente de Lucid-Lab, "
    "Version 1.0, accessibles à https://lucid-lab.fr/cgv.",
    space_after=177800,
)

# ── Prochaines étapes ──────────────────────────────────────────────────────────
add_h1(doc, "Prochaines étapes")
for item in [
    "Planifier un appel de cadrage pour valider définitivement le périmètre technique, les priorités et les premiers jalons avant démarrage.",
    "Signer la présente Proposition valant Bon de commande et procéder au premier versement (30 %, soit 3 240,00 € TTC).",
    "Planifier le kick-off de démarrage.",
    "Confirmer le panel de clients B2B à onboarder en Mois 3.",
]:
    add_bullet(doc, item)

# ── Cadre contractuel ──────────────────────────────────────────────────────────
add_h1(doc, "Cadre contractuel et démarrage")
add_body(doc,
    "Démarrage des Prestations : après signature du Bon de Commande accompagnant la présente Proposition "
    "et réception effective du premier paiement sur le compte bancaire de Lucid-Lab "
    "(IBAN FR76 1732 8844 0043 2662 8862 178, BIC SWNBFR22).",
    space_after=50800,
)
add_body(doc,
    "Documents contractuels remis : (i) la présente Proposition, (ii) le Bon de Commande à signer, "
    "(iii) le Contrat de Prestation, (iv) les CGV v1.0, (v) le cas échéant l'Accord de sous-traitance des données (DPA).",
    space_after=50800,
)
add_body(doc,
    "Couverture assurantielle : Lucid-Lab est assurée en Responsabilité Civile Professionnelle auprès de Hiscox SA "
    "(police RCPLP 3695175450, plafond 100 000 € par période d'assurance, monde entier hors USA/Canada).",
    space_after=177800,
)

# ── Acceptation ────────────────────────────────────────────────────────────────
add_h1(doc, "Acceptation valant commande ferme")
add_body(doc, "Le Client déclare avoir pris connaissance et accepter sans réserve :", space_after=50800)
for item in [
    "le périmètre, les livrables et le calendrier décrits ci-dessus ;",
    "la modalité tarifaire retenue et l'échéancier de paiement associé ;",
    "les Conditions Générales de Vente Lucid-Lab v1.0 accessibles sur lucid-lab.fr/cgv ;",
    "l'obligation de versement préalable au démarrage des Prestations (art. 5 CGV).",
]:
    add_bullet(doc, item)

add_body(doc,
    "La signature de la présente Proposition vaut acceptation de l'offre, formation du Contrat et commande ferme, "
    "et déclenche l'émission d'une facture pro forma puis l'exécution des Prestations à compter de la réception "
    "effective du premier paiement.",
    space_after=101600,
)

p_modal = doc.add_paragraph()
_run(p_modal, "Modalité retenue : ☑ Forfait Phase 1 B2B, 9 000,00 € HT / 3 mois, échéancier 30/40/30.", bold=True)
_set_spacing(p_modal, after_emu=76200)

p_mention = doc.add_paragraph()
_run(p_mention,
    "Mention manuscrite obligatoire : « Lu et approuvé, bon pour accord et commande ferme, "
    "Prix : 10 800,00 € TTC, Modalité : Forfait Phase 1 B2B ».", bold=True)
_set_spacing(p_mention, after_emu=177800)

# ── Signature ──────────────────────────────────────────────────────────────────
add_signature_table(doc, "Sinibaldi Agency / Clément Sinibaldi")

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = OUT / FILENAME
doc.save(str(out_path))
print(f"✓ Sauvegardé : {out_path}")
