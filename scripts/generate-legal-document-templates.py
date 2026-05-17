#!/usr/bin/env python3
"""
Generate Lucid-Lab legal document templates from the source Word models.

The source models remain untouched. This script only adds the Lucid-Lab logo
next to the existing header wordmark and writes generated copies.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "docs" / "legal-templates" / "source"
OUTPUT_DIR = ROOT / "docs" / "legal-templates" / "generated"
LOGO_PATH = ROOT / "public" / "logo.png"
LOGO_WIDTH = Cm(0.72)

TEMPLATES = [
    {
        "source": "03_Proposition-valant-Bon-de-commande-modele.docx",
        "output": "03_Proposition-valant-Bon-de-commande-modele-logo.docx",
        "purpose": "Proposition commerciale valant bon de commande",
    },
    {
        "source": "04_Contrat-prestation-modele.docx",
        "output": "04_Contrat-prestation-modele-logo.docx",
        "purpose": "Contrat de prestation de services",
    },
]


def insert_logo_before_header_wordmark(doc: Document) -> int:
    """Insert the logo before header paragraphs beginning with Lucid-Lab."""
    inserted = 0

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if not paragraph.text.strip().startswith("Lucid-Lab"):
                continue

            logo_run = paragraph.add_run()
            logo_run.add_picture(str(LOGO_PATH), width=LOGO_WIDTH)

            run_element = logo_run._r
            paragraph._p.remove(run_element)
            insert_index = 1 if paragraph._p.pPr is not None else 0
            paragraph._p.insert(insert_index, run_element)
            inserted += 1

    return inserted


def generate_template(config: dict[str, str]) -> Path:
    source_path = SOURCE_DIR / config["source"]
    output_path = OUTPUT_DIR / config["output"]

    if not source_path.exists():
        raise FileNotFoundError(f"Missing source model: {source_path}")
    if not LOGO_PATH.exists():
        raise FileNotFoundError(f"Missing logo: {LOGO_PATH}")

    doc = Document(source_path)
    logo_count = insert_logo_before_header_wordmark(doc)
    if logo_count == 0:
        raise RuntimeError(f"No Lucid-Lab header found in {source_path.name}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Generated {output_path.relative_to(ROOT)} — {config['purpose']} ({logo_count} logo)")
    return output_path


def main() -> int:
    for config in TEMPLATES:
        generate_template(config)
    return 0


if __name__ == "__main__":
    sys.exit(main())