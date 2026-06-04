#!/usr/bin/env python3
"""
Generate filled BDC + Contrat for two clients:
- Clément Sinibaldi (SASU)
- Le Jardin d'Eden (Crèche Eden)
"""
from __future__ import annotations
import copy
import shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SOURCE_BDC = ROOT / "docs/legal-templates/source/03_Proposition-valant-Bon-de-commande-modele.docx"
SOURCE_CTR = ROOT / "docs/legal-templates/source/04_Contrat-prestation-modele.docx"
SOURCE_BDC_LOGO = ROOT / "docs/legal-templates/source/03_Proposition-valant-Bon-de-commande-modele-logo.docx"
SOURCE_CTR_LOGO = ROOT / "docs/legal-templates/source/04_Contrat-prestation-modele-logo.docx"
OUT = ROOT / "docs/legal-templates/generated"
OUT.mkdir(parents=True, exist_ok=True)

# ─── Client definitions ──────────────────────────────────────────────────────

CLIENTS = [
    {
        "id": "sophia-kanouni-bsp37",
        "logo": True,
        "bdc_ref": "PROP-2026-009",
        "ctr_ref": "CTR-2026-009",
        "bc_ref": "BC-2026-009",
        "date": "23/05/2026",
        "client_block_bdc": (
            "Le Client — Sophia Kanouni, agissant pour le compte de BSP 37, "
            "groupement familial dont le siège social est situé Route du Coutureau, "
            "\"Le Coudray\", 37120 Lémeré, France, numéro SIRET [À COMPLÉTER], "
            "représentée par Sophia Kanouni."
        ),
        "client_block_ctr": (
            "Sophia Kanouni, agissant pour le compte de BSP 37, "
            "dont le siège social est situé Route du Coutureau, \"Le Coudray\", 37120 Lémeré, France, "
            "numéro SIRET [À COMPLÉTER], représentée par Sophia Kanouni,"
        ),
        "interlocteur": "Sophia Kanouni — BSP 37",
        "lieu": "Lémeré",
        "mission_title": "Refonte + maintenance du site web bsp37.com et gestion Meta Ads à partir du 1er septembre 2026.",
        "mission_description": (
            "Refonte complète du site web bsp37.com avec maintenance mensuelle, puis gestion des campagnes "
            "Meta Ads pour l'activité film solaire à partir du 1er septembre 2026, incluant la création "
            "de deux publicités de lancement. Le forfait mensuel est de 250,00 € HT, composé de 150,00 € HT "
            "pour le site web et 100,00 € HT pour la gestion Meta Ads."
        ),
        "livrable_1": "Site web bsp37.com refondu sur technologie moderne : catalogue, galerie, services, formulaires, déploiement.",
        "livrable_2": "Maintenance mensuelle du site : mises à jour, corrections, suivi SEO et performances.",
        "livrable_3": "Gestion Meta Ads à partir du 1er septembre 2026 + création de deux publicités de lancement.",
        "calendrier": "Démarrage à compter de la signature et réception du premier paiement.",
        "prix_ht": "3 250,00 € (250,00 € création de deux publicités + 12 × 250,00 €, dont 150,00 € site web et 100,00 € gestion Meta Ads mensuels)",
        "tva": "650,00 €",
        "prix_ttc": "3 900,00 €",
        "modalite": "☑ Mensuel 12 mois — site web + gestion Meta Ads à partir du 1er septembre 2026 + frais unique création de deux publicités",
        "echeancier": "300,00 € TTC (création de deux publicités, frais unique) + 300,00 € TTC × 12 mois (180,00 € TTC site web + 120,00 € TTC gestion Meta Ads)",
        "ref_client_rib": "BC-2026-009 / BSP 37",
        "prop_ref_ctr": "PROP-2026-009 — BSP 37 — 23/05/2026",
        "perimetre_ctr": (
            "refonte complète du site web bsp37.com (audit du site Wix existant, développement sur technologie moderne, "
            "intégration du catalogue produits, galerie de réalisations, pages de services et formulaires de contact, "
            "déploiement et configuration du nom de domaine) ; "
            "maintenance mensuelle du site (mises à jour, corrections, suivi des performances SEO) ; "
            "gestion des campagnes Meta Ads pour l'activité film solaire à partir du 1er septembre 2026, "
            "incluant optimisation, tests visuels, reporting mensuel et création de deux publicités de lancement. "
            "La gestion Meta Ads est facturée 100,00 € HT par mois et la création des deux publicités est facturée "
            "250,00 € HT en frais unique."
        ),
        "hors_perimetre": (
            "toute prestation non explicitement listée dans la Proposition Commerciale, notamment le budget "
            "publicitaire Meta Ads, la production de publicités ou contenus supplémentaires au-delà des deux publicités "
            "prévues, toute intégration supplémentaire non prévue et les coûts d'hébergement ou d'abonnements tiers."
        ),
        "prix_ht_ctr": "3 250,00 €",
        "tva_ctr": "650,00 €",
        "prix_ttc_ctr": "3 900,00 €",
        "mensualite_ttc": "300,00 €",
    },
    {
        "id": "clement-sinibaldi",
        "bdc_ref": "PROP-2026-007",
        "ctr_ref": "CTR-2026-007",
        "bc_ref": "BC-2026-007",
        "date": "18/05/2026",
        "client_block_bdc": (
            "Le Client — Clément Sinibaldi, SASU, siège social : 3 avenue Gozza, "
            "83000 Toulon, France, immatriculée sous le numéro SIRET 947 876 892 00011, "
            "représentée par Clément Sinibaldi en qualité de Président."
        ),
        "client_block_ctr": (
            "Clément Sinibaldi, SASU, dont le siège social est situé 3 avenue Gozza, "
            "83000 Toulon, France, immatriculée sous le numéro SIRET 947 876 892 00011, "
            "représentée par Clément Sinibaldi en qualité de Président,"
        ),
        "interlocteur": "Clément Sinibaldi — Président — contact@sinibaldi-architecte.fr",
        "lieu": "Toulon",
        "mission_title": "Nouvelle page WordPress, acquisition Meta Ads et agent IA d'inspiration Instagram pour développer le service de conseil.",
        "mission_description": (
            "Développement d'une page dédiée au service de conseil intégrée au site WordPress existant, "
            "mise en place d'une automatisation des posts blog, gestion des campagnes Meta Ads pour la "
            "génération de leads, et déploiement d'un agent IA proposant des idées de contenu Instagram."
        ),
        "livrable_1": "Page WordPress dédiée au service de conseil, intégrée au site existant + automatisation des posts blog.",
        "livrable_2": "Gestion campagnes Meta Ads (génération de leads et acquisition).",
        "livrable_3": "Agent IA : trames inspirationnelles mensuelles Instagram, suggestions de posts et améliorations du contenu existant.",
        "calendrier": "Démarrage à compter de la signature et réception du premier paiement.",
        "prix_ht": "1 908,00 € (12 × 159,00 €)",
        "tva": "381,60 €",
        "prix_ttc": "2 289,60 €",
        "modalite": "☐ One-shot — 1 590,00 € HT / 1 908,00 € TTC en une fois avant démarrage\n☑ Mensuel 12 mois — engagement ferme",
        "echeancier": "190,80 € TTC × 12 mois  OU  1 908,00 € TTC en une fois",
        "ref_client_rib": "BC-2026-007 / Sinibaldi Architecte",
        "prop_ref_ctr": "PROP-2026-007 — Sinibaldi Architecte — 14/05/2026",
        "perimetre_ctr": (
            "Développement d'une page WordPress dédiée au service de conseil, intégrée au site existant ; "
            "mise en place d'une automatisation des posts blog ; gestion des campagnes Meta Ads "
            "(objectif : génération de leads et acquisition) ; développement et déploiement d'un agent IA "
            "proposant des idées de contenu Instagram (trames mensuelles, suggestions de posts, améliorations du contenu existant)."
        ),
        "hors_perimetre": (
            "toute prestation non explicitement listée dans la Proposition Commerciale, notamment la refonte "
            "graphique globale du site, la conception de nouveaux services autres que le service de conseil, "
            "les intégrations supplémentaires non prévues, et le sujet SaaS (à traiter dans un avenant séparé)."
        ),
        "prix_ht_ctr": "1 908,00 €",
        "tva_ctr": "381,60 €",
        "prix_ttc_ctr": "2 289,60 €",
        "mensualite_ttc": "190,80 €",
    },
    {
        "id": "eden-creche",
        "bdc_ref": "PROP-2026-008",
        "ctr_ref": "CTR-2026-008",
        "bc_ref": "BC-2026-008",
        "date": "18/05/2026",
        "client_block_bdc": (
            "Le Client — Le Jardin d'Eden, dont le siège social est situé Rue des Deux Eglises 14-16, "
            "1000 Bruxelles, Belgique, immatriculée sous le numéro BCE BE0668.411.657, "
            "représentée par son représentant légal dûment habilité."
        ),
        "client_block_ctr": (
            "Le Jardin d'Eden, dont le siège social est situé Rue des Deux Eglises 14-16, "
            "1000 Bruxelles, Belgique, immatriculée sous le numéro BCE BE0668.411.657, "
            "représentée par son représentant légal dûment habilité,"
        ),
        "interlocteur": "À compléter — Le Jardin d'Eden",
        "lieu": "Bruxelles",
        "mission_title": "Automatisation administrative, acquisition Meta Ads et socle SEO pour lisser l'occupation de la crèche.",
        "mission_description": (
            "Mise en place d'un back-office Google Sheets automatisé remplaçant le fichier Excel principal, "
            "gestion documentaire automatisée avec relances et classement Drive, gestion des campagnes Meta Ads, "
            "ajout WhatsApp Business et programmation mensuelle de posts blog SEO."
        ),
        "livrable_1": "Google Sheets avec formules, automatisations et suivi virements ; connexion Drive et emails entrants.",
        "livrable_2": "Gestion Meta Ads + WhatsApp Business ; classement auto des documents et relances email aux échéances.",
        "livrable_3": "Routine SEO : programmation mensuelle de posts blog orientés référencement.",
        "calendrier": "Démarrage à compter de la signature et réception du premier paiement.",
        "prix_ht": "3 108,00 € (12 × 259,00 €)",
        "tva": "621,60 €",
        "prix_ttc": "3 729,60 €",
        "modalite": "☐ One-shot — 2 590,00 € HT / 3 108,00 € TTC en une fois avant démarrage\n☑ Mensuel 12 mois — engagement ferme",
        "echeancier": "310,80 € TTC × 12 mois  OU  3 108,00 € TTC en une fois",
        "ref_client_rib": "BC-2026-008 / Le Jardin d'Eden",
        "prop_ref_ctr": "PROP-2026-008 — Crèche Eden — 14/05/2026",
        "perimetre_ctr": (
            "Transposition du fichier Excel principal vers Google Sheets avec formules, automatisations et "
            "suivi des virements par client ; connexion entre le fichier principal, Google Drive et les emails entrants ; "
            "gestion des disponibilités ; classement automatique des documents reçus, détection des pièces manquantes ; "
            "relances automatiques par email aux échéances utiles ; demandes automatiques de signature et de documentation ; "
            "automatisation des demandes d'avis Google ; prise en charge de la gestion Meta Ads ; "
            "ajout d'un WhatsApp Business ; programmation mensuelle de posts blog orientés SEO."
        ),
        "hors_perimetre": (
            "toute prestation non explicitement listée dans la Proposition Commerciale, notamment la refonte "
            "complète de l'infrastructure IT, la création d'un site web, et toute automatisation "
            "supplémentaire non décrite dans la proposition."
        ),
        "prix_ht_ctr": "3 108,00 €",
        "tva_ctr": "621,60 €",
        "prix_ttc_ctr": "3 729,60 €",
        "mensualite_ttc": "310,80 €",
    },
]

# ─── Text replacement utilities ──────────────────────────────────────────────

def replace_in_para(para, replacements: dict[str, str]):
    """Replace in runs, falling back to paragraph rebuild for cross-run placeholders."""
    for old, new in replacements.items():
        if old not in para.text:
            continue
        # Try single-run replacement first (preserves formatting)
        replaced = False
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                replaced = True
                break
        if not replaced and para.runs:
            # Cross-run case: rebuild from first run
            full = para.text.replace(old, new)
            para.runs[0].text = full
            for run in para.runs[1:]:
                run.text = ""

def replace_in_doc(doc: Document, replacements: dict[str, str]):
    # Body paragraphs
    for para in doc.paragraphs:
        replace_in_para(para, replacements)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para, replacements)
    # Headers / footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_para(para, replacements)
        for para in section.footer.paragraphs:
            replace_in_para(para, replacements)


# ─── Document generators ──────────────────────────────────────────────────────

def build_bdc(client: dict) -> Path:
    src = SOURCE_BDC_LOGO if client.get("logo") else SOURCE_BDC
    doc = Document(str(src))
    replacements = {
        "PROP-[NUMÉRO]": client["bdc_ref"],
        "[JJ/MM/AAAA]": client["date"],
        "Le Client — [Raison sociale], [forme juridique] au capital de [montant] €, siège social : [adresse], immatriculée au RCS de [ville] sous le numéro [SIREN], numéro de TVA intracommunautaire [FR XX XXXXXXXXX], représentée par [Prénom NOM] en qualité de [fonction].": client["client_block_bdc"],
        "[Titre de la mission]": client["mission_title"],
        "[Décrire en 2-3 lignes le périmètre principal — agents IA, automatisations, développement web, gestion ads, contenus]": client["mission_description"],
        "•  [Livrable 1]": f"•  {client['livrable_1']}",
        "•  [Livrable 2]": f"•  {client['livrable_2']}",
        "•  [Livrable 3]": f"•  {client['livrable_3']}",
        "démarrage à compter du [JJ/MM/AAAA] sous réserve de la signature du présent Bon de Commande et de la réception du premier paiement": client["calendrier"],
        "[Montant] €\n[Montant] €\n[Montant] €": "",  # handled via table below
        "BC-[NUMÉRO] / [Nom du Client]": client["ref_client_rib"],
        "[Total TTC]": client["prix_ttc"].split(" ")[0] if " " in client["prix_ttc"] else client["prix_ttc"],
        "[one-shot / mensuel 12 mois]": "one-shot / mensuel 12 mois",
    }
    replace_in_doc(doc, replacements)

    # Table replacements
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if cells[0] == "Prix HT total":
                row.cells[1].paragraphs[0].runs[0].text = client["prix_ht"]
            elif cells[0] == "TVA (20 %)":
                row.cells[1].paragraphs[0].runs[0].text = client["tva"]
            elif cells[0] == "Prix TTC total":
                row.cells[1].paragraphs[0].runs[0].text = client["prix_ttc"]
            elif cells[0] == "Modalité retenue":
                row.cells[1].paragraphs[0].runs[0].text = client["modalite"]
            elif cells[0] == "Échéancier":
                row.cells[1].paragraphs[0].runs[0].text = client["echeancier"]
            elif cells[0] == "Référence à indiquer":
                row.cells[1].paragraphs[0].runs[0].text = client["ref_client_rib"]
            elif "Lieu :" in cells[0]:
                row.cells[0].paragraphs[0].runs[0].text = f"Lieu : {client['lieu']}"

    out = OUT / f"{client['id']}-bon-de-commande.docx"
    doc.save(str(out))
    return out


def build_ctr(client: dict) -> Path:
    src = SOURCE_CTR_LOGO if client.get("logo") else SOURCE_CTR
    doc = Document(str(src))
    replacements = {
        "CTR-[NUMÉRO]": client["ctr_ref"],
        "[JJ/MM/AAAA]": client["date"],
        "[Raison sociale], [forme juridique] au capital de [montant] euros, dont le siège social est situé [adresse complète], immatriculée au RCS de [ville] sous le numéro [SIREN], numéro de TVA intracommunautaire [FR XX XXXXXXXXX], représentée par [Prénom NOM] en qualité de [fonction],": client["client_block_ctr"],
        "[Réf. proposition]": client["prop_ref_ctr"],
        "Périmètre des Prestations : [Description détaillée — agents IA, bots conversationnels, automatisations métier, développement web, gestion campagnes Meta Ads, production de contenus SEO, etc.]": f"Périmètre des Prestations : {client['perimetre_ctr']}",
        "Hors périmètre : toute prestation non explicitement listée dans la Proposition Commerciale, notamment [exemples : refonte graphique du site, conception de nouveaux services, intégrations supplémentaires non prévues].": f"Hors périmètre : {client['hors_perimetre']}",
        "Date de démarrage envisagée : [JJ/MM/AAAA].": f"Date de démarrage envisagée : dès réception des accès et du premier paiement.",
        "Prix total HT : [Montant] €. TVA 20 % : [Montant] €. Prix TTC : [Montant] €.": f"Prix total HT : {client['prix_ht_ctr']}. TVA 20 % : {client['tva_ctr']}. Prix TTC : {client['prix_ttc_ctr']}.",
        "[Montant TTC] €": f"{client['mensualite_ttc']} €",
        "[Prénom NOM — fonction — email — téléphone]": client["interlocteur"],
    }
    replace_in_doc(doc, replacements)

    out = OUT / f"{client['id']}-contrat.docx"
    doc.save(str(out))
    return out


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    for client in CLIENTS:
        bdc = build_bdc(client)
        ctr = build_ctr(client)
        print(f"✓ {client['id']}")
        print(f"  BDC  → {bdc.relative_to(ROOT)}")
        print(f"  CTR  → {ctr.relative_to(ROOT)}")

if __name__ == "__main__":
    main()
